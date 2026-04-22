"""DOCX generator tests: verify the produced bytes are a valid Word file,
preserve the section headings from the template, and swap bullets to content
pulled from the truth model.
"""

import io

import pytest
from docx import Document

from app.services.data_loader import load_truth_model
from app.services.resume_docx import generate_tailored_resume_bytes


ARCHETYPES = [
    "A_general_ai_platform",
    "B_fintech_transaction_systems",
    "C_data_streaming_systems",
    "D_distributed_systems",
]

JD_FINTECH = (
    "Senior Backend Engineer, Payments Platform. Transaction integrity, "
    "auditability, compliance, entitlements, Kafka, distributed services."
)


def _paragraphs_text(doc):
    return [p.text for p in doc.paragraphs]


@pytest.mark.parametrize("archetype", ARCHETYPES)
def test_generator_returns_valid_docx(archetype):
    blob = generate_tailored_resume_bytes(archetype, JD_FINTECH)
    assert blob[:2] == b"PK", "DOCX files are zip archives starting with 'PK'"
    doc = Document(io.BytesIO(blob))
    texts = _paragraphs_text(doc)
    joined = "\n".join(texts).lower()
    assert "professional experience" in joined
    assert "education" in joined


def test_bullets_replaced_come_from_truth_model():
    """Every replaced bullet should be either a verbatim core_fact from the
    truth model or an unchanged template bullet we did not touch. Any text
    that matches a role header or heading is left alone."""
    blob = generate_tailored_resume_bytes("B_fintech_transaction_systems", JD_FINTECH)
    doc = Document(io.BytesIO(blob))

    truth = load_truth_model()
    facts = set()
    for role in truth.get("roles", []):
        facts.update(role.get("core_facts", []))

    # At least one paragraph (post-experience, pre-education) should be a
    # verbatim truth-model core_fact after we pulled them in.
    texts = _paragraphs_text(doc)
    lower = [t.lower() for t in texts]
    exp_idx = next(i for i, t in enumerate(lower) if "professional experience" in t)
    edu_idx = next(i for i, t in enumerate(lower) if t.strip().startswith(">") and "education" in t)

    body_texts = [t.strip() for t in texts[exp_idx + 1 : edu_idx] if t.strip()]
    assert body_texts, "expected some body paragraphs in the experience section"

    matches = sum(1 for t in body_texts if t in facts)
    assert matches >= 1, (
        f"expected at least one verbatim truth-model bullet in the experience "
        f"section; body had {len(body_texts)} lines, matched {matches}"
    )


def test_current_role_gets_at_least_four_bullets_in_generated_docx():
    """Template A ships only 2 Citibank bullet slots, but the generator should
    expand that to >= 4 bullets for the current role so the resume's anchor
    section isn't thin."""
    blob = generate_tailored_resume_bytes("A_general_ai_platform", JD_FINTECH)
    doc = Document(io.BytesIO(blob))
    texts = _paragraphs_text(doc)

    exp_idx = next(i for i, t in enumerate(texts) if "professional experience" in t.lower())
    # Find the Citibank header (the current role).
    citi_idx = next(
        i for i, t in enumerate(texts)
        if i > exp_idx and "citibank" in t.lower() and "|" in t
    )
    # Walk to the next role header (or Education).
    next_role_idx = next(
        (i for i, t in enumerate(texts) if i > citi_idx and ("|" in t and "@" not in t and i > citi_idx + 1) or t.strip().startswith(">EDUCATION")),
        len(texts),
    )
    body = [t.strip() for t in texts[citi_idx + 1 : next_role_idx] if t.strip()]
    # Filter out anything that looks like a role header (contains `|`).
    bullets = [t for t in body if "|" not in t]
    assert len(bullets) >= 4, (
        f"expected >= 4 Citibank bullets in generated A-archetype resume, got {len(bullets)}: {bullets}"
    )


def test_generator_normalizes_stray_bullet_glyphs():
    """Template A's Education section uses `●` on the B.S. line. The generator
    should normalize it to `•` so the resume uses a consistent glyph."""
    blob = generate_tailored_resume_bytes("A_general_ai_platform", JD_FINTECH)
    doc = Document(io.BytesIO(blob))
    texts = _paragraphs_text(doc)
    joined = "\n".join(texts)
    assert "\u25cf" not in joined, "Heavy bullet glyph ● should have been normalized to •"


def test_summary_is_tailored_and_nonempty():
    blob = generate_tailored_resume_bytes("C_data_streaming_systems", JD_FINTECH)
    doc = Document(io.BytesIO(blob))
    texts = [p.text for p in doc.paragraphs]
    lower = [t.lower() for t in texts]
    lang_idx = next(i for i, t in enumerate(lower) if "languages and technologies" in t)
    pre = [t.strip() for t in texts[:lang_idx] if t.strip()]
    assert pre, "summary section should have text before LANGUAGES heading"
    # Last non-empty paragraph before LANGUAGES is the summary body and should
    # be substantive (not just the title heading starting with '>').
    candidates = [t for t in pre if not t.startswith(">")]
    assert candidates, "expected a body summary paragraph"
    assert any(len(t) > 60 for t in candidates)
