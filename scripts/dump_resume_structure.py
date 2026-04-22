"""One-off: dump the paragraph/run structure of a .docx so we can design a
template-aware tailoring pipeline."""

import io
import sys
from pathlib import Path

from docx import Document

# Force ascii-safe output on Windows consoles that default to cp1252.
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")


def _safe(s: str) -> str:
    return s.replace("\n", "\\n")


def dump(path: Path) -> None:
    doc = Document(path)
    print(f"\n========== {path.name} ==========")
    print(f"sections: {len(doc.sections)}  tables: {len(doc.tables)}")
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else "-"
        text = _safe(p.text)
        if not text.strip() and not p.runs:
            continue
        print(f"  [{i:03d}] style={style!r:<30} text={text[:110]!r}")
    for ti, table in enumerate(doc.tables):
        print(f"  TABLE {ti}: rows={len(table.rows)} cols={len(table.columns)}")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    txt = _safe(p.text)
                    if txt.strip():
                        print(f"    t{ti} r{ri} c{ci} p{pi}: {txt[:110]!r}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1] / "data" / "source_resumes"
    targets = sys.argv[1:] or [p.name for p in sorted(root.iterdir()) if p.suffix == ".docx"]
    for t in targets:
        dump(root / t)
