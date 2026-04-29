"""Template-aware DOCX resume generator.

Opens the archetype's source resume as a template and rewrites ONLY the
summary paragraph(s) and the bullet paragraphs under each role, preserving:
- the contact table
- section headings (SUMMARY, LANGUAGES AND TECHNOLOGIES, PROFESSIONAL EXPERIENCE, EDUCATION)
- languages / skills content
- role header lines (Company | Title | Location ... date)
- education section
- all runs, fonts, colors, bullet glyphs, and styles

Guardrail: the text that replaces bullets is always pulled from
`master_truth_model.json` via `resume_tailor.rank_role_bullets`. Nothing is
invented.
"""

from __future__ import annotations

import copy
import io
import re
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.text.paragraph import Paragraph

from app.services.data_loader import load_archetypes, load_truth_model
from app.services.resume_tailor import (
    _normalize_company,
    draft_summary,
    rank_role_bullets,
)

try:  # LLM is optional; import lazily in the function that uses it.
    from app.services.llm_rewrite import (
        is_available as _llm_available,
        rewrite_bullets as _llm_rewrite_bullets,
        rewrite_summary as _llm_rewrite_summary,
    )
except Exception:  # pragma: no cover
    _llm_available = lambda: False  # type: ignore
    _llm_rewrite_bullets = lambda bullets, jd: bullets  # type: ignore
    _llm_rewrite_summary = lambda s, jd, a: s  # type: ignore

BASE_DIR = Path(__file__).resolve().parents[2]
SOURCE_RESUMES_DIR = BASE_DIR / "data" / "source_resumes"

# Match things like "●", "•", "▪", or "o " at start of text.
_LEADING_BULLET_RE = re.compile(r"^[\s\u00a0]*([\u2022\u25cf\u25a0\u25aa\u2023\u2043\u25e6\-o])\s*")

# Heading markers (case-insensitive "contains").
_MARKER_LANGUAGES = "languages and technologies"
_MARKER_EXPERIENCE = "professional experience"
_MARKER_EDUCATION = "education"


# ---------- paragraph text manipulation ----------


def _paragraph_is_blank(p: Paragraph) -> bool:
    return not p.text.strip()


def _paragraph_text_lc(p: Paragraph) -> str:
    return (p.text or "").strip().lower()


def _has_marker(p: Paragraph, marker: str) -> bool:
    return marker in _paragraph_text_lc(p)


def _strip_leading_bullet(text: str) -> tuple[str, str]:
    """Return (prefix, rest) where prefix is the bullet glyph + surrounding
    whitespace if present, else empty string."""
    m = _LEADING_BULLET_RE.match(text)
    if not m:
        return "", text
    return text[: m.end()], text[m.end() :]


_CANONICAL_BULLET_PREFIX = "\u2022 "  # "• " — the dominant glyph across templates.


def _normalize_bullet_prefix(prefix: str) -> str:
    """Normalize any detected bullet glyph to a single `• ` so the generated
    resume doesn't mix `•`, `●`, and `▪`. If the original had no glyph (e.g. a
    heading-styled bullet paragraph), return empty string — we don't want to
    force a glyph where the template never had one."""
    if not prefix:
        return ""
    return _CANONICAL_BULLET_PREFIX


def _replace_paragraph_text(p: Paragraph, new_text: str) -> None:
    """Replace the paragraph's visible text while keeping the first run's
    formatting. Normalizes a leading bullet glyph if the paragraph had one."""
    original = p.text or ""
    prefix, _rest = _strip_leading_bullet(original)
    final_text = _normalize_bullet_prefix(prefix) + new_text.strip()

    runs = list(p.runs)
    if not runs:
        p.add_run(final_text)
        return

    for extra in runs[1:]:
        extra.text = ""
    runs[0].text = final_text


def _clone_paragraph_after(source_p: Paragraph, text: str) -> Paragraph:
    """Deep-copy `source_p`'s XML element, place it immediately after the
    source, and rewrite its text to `text`. This preserves the source
    paragraph's style, indentation, font runs, and bullet formatting — which
    is how we add extra bullet slots to a template that was short on them.
    Returns the new Paragraph wrapper around the inserted element."""
    new_el = copy.deepcopy(source_p._element)
    source_p._element.addnext(new_el)
    new_para = Paragraph(new_el, source_p._parent)
    _replace_paragraph_text(new_para, text)
    return new_para


# ---------- document structure detection ----------


def _find_indices(paragraphs: List[Paragraph]) -> dict:
    """Locate key section indices. Missing sections map to None."""
    idx = {"languages": None, "experience": None, "education": None}
    for i, p in enumerate(paragraphs):
        t = _paragraph_text_lc(p)
        if idx["languages"] is None and _MARKER_LANGUAGES in t:
            idx["languages"] = i
        elif idx["experience"] is None and _MARKER_EXPERIENCE in t:
            idx["experience"] = i
        elif idx["education"] is None and _MARKER_EDUCATION in t and t.startswith(">"):
            idx["education"] = i
    return idx


def _is_role_header(p: Paragraph) -> bool:
    """A role header looks like 'Company | Title | Location ... date'."""
    t = (p.text or "").strip()
    if "|" not in t:
        return False
    if len(t) > 300:
        return False
    # Heuristic: at least two '|' separators OR explicit location/date pattern.
    return t.count("|") >= 1 and "@" not in t


def _extract_company_and_title(text: str) -> tuple[str, Optional[str]]:
    """Given a role-header line like 'Company | Title | Location ... date',
    return (company, title) — title may be None when not separable."""
    parts = [p.strip() for p in (text or "").split("|")]
    company = parts[0] if parts else ""
    title = parts[1] if len(parts) >= 2 else None
    if title:
        # Strip trailing tabs/dates that sometimes hug the title (e.g. "Senior Backend Engineer ...").
        title = re.split(r"\s{2,}|\t", title)[0].strip() or None
    return company, title


# ---------- main entry point ----------


def _resolve_template_path(archetype_id: str) -> Path:
    archetypes = load_archetypes()
    meta = archetypes.get(archetype_id) or archetypes.get("A_general_ai_platform") or {}
    filename = meta.get("source_resume")

    candidates: List[Path] = []
    if filename:
        candidates.append(SOURCE_RESUMES_DIR / filename)

    prefix_map = {
        "A_general_ai_platform": "MM_Resume_4_9_26_A",
        "B_fintech_transaction_systems": "MM_Resume_4_9_26_B",
        "C_data_streaming_systems": "MM_Resume_4_9_26_C",
        "D_distributed_systems": "MM_Resume_4_9_26_D",
        "E_staff_backend": "MM_Resume_4_9_26_E",
        "E_core_backend": "MM_Resume_4_9_26_E",
        "E_resume": "MM_Resume_4_9_26_E",
    }
    prefix = prefix_map.get(archetype_id)
    if prefix:
        # Fall back to the first file matching the archetype's letter.
        for p in sorted(SOURCE_RESUMES_DIR.glob(f"{prefix}*.docx")):
            candidates.append(p)

    for c in candidates:
        if c.exists():
            return c

    raise FileNotFoundError(
        f"No DOCX template found for archetype {archetype_id}. "
        f"Looked in {SOURCE_RESUMES_DIR}."
    )


def _rewrite_summary(paragraphs: List[Paragraph], summary_text: str, idx: dict) -> None:
    languages_idx = idx.get("languages")
    if languages_idx is None:
        return
    # The summary body is the first 1-2 non-blank paragraphs before `>LANGUAGES`
    # that are not the title heading (headings starting with `>`).
    summary_candidates: List[int] = []
    for i in range(languages_idx - 1, -1, -1):
        p = paragraphs[i]
        text = (p.text or "").strip()
        if not text:
            continue
        if text.startswith(">"):
            break
        summary_candidates.append(i)
    summary_candidates.reverse()

    if not summary_candidates:
        return

    # Replace the FIRST summary paragraph with the full tailored summary.
    _replace_paragraph_text(paragraphs[summary_candidates[0]], summary_text)
    # Blank any additional summary paragraphs so we don't leave stale text.
    for i in summary_candidates[1:]:
        _replace_paragraph_text(paragraphs[i], "")


_MIN_CURRENT_ROLE_BULLETS = 4


def _is_current_role_company(company: str) -> bool:
    """True when this company corresponds to the truth model's `is_current` role.
    Used to give the current role extra bullet slots when the template was thin
    (e.g., archetype A ships only 2 Citibank bullets, but we always want >= 4)."""
    normalized = _normalize_company(company)
    for role in load_truth_model().get("roles", []):
        if not role.get("is_current"):
            continue
        if _normalize_company(role.get("company", "")) == normalized:
            return True
    return False


def _rewrite_experience(
    paragraphs: List[Paragraph],
    idx: dict,
    job_description: str,
    use_llm: bool = False,
) -> int:
    """Return the number of roles whose bullets were rewritten.

    When the truth model has more high-scoring bullets for the current role
    than the template provides slots for, we clone the last bullet paragraph
    (preserving its style) and fill in the extra text — so template A's 2
    Citibank slots grow to 4+ and match templates B/D.
    """
    exp_idx = idx.get("experience")
    end_idx = idx.get("education") or len(paragraphs)
    if exp_idx is None:
        return 0

    truth_companies = {
        _normalize_company(r["company"]) for r in load_truth_model().get("roles", [])
    }

    rewritten_count = 0

    i = exp_idx + 1
    current_company: Optional[str] = None
    current_title: Optional[str] = None
    current_bullet_indices: List[int] = []

    def _flush():
        nonlocal current_company, current_title, current_bullet_indices, rewritten_count
        if current_company and current_bullet_indices:
            normalized = _normalize_company(current_company)
            if normalized in truth_companies:
                is_current = _is_current_role_company(current_company)
                target_slot_count = len(current_bullet_indices)
                if is_current and target_slot_count < _MIN_CURRENT_ROLE_BULLETS:
                    target_slot_count = _MIN_CURRENT_ROLE_BULLETS

                bullets = rank_role_bullets(
                    job_description,
                    company=current_company,
                    title=current_title,
                    limit=max(target_slot_count, len(current_bullet_indices)),
                )
                if use_llm and _llm_available() and bullets:
                    bullets = _llm_rewrite_bullets(bullets, job_description)

                for slot, bullet_idx in enumerate(current_bullet_indices):
                    if slot < len(bullets):
                        _replace_paragraph_text(paragraphs[bullet_idx], bullets[slot])

                extras_needed = max(0, target_slot_count - len(current_bullet_indices))
                if extras_needed > 0 and current_bullet_indices and len(bullets) > len(current_bullet_indices):
                    last_bullet_p = paragraphs[current_bullet_indices[-1]]
                    extras_available = bullets[len(current_bullet_indices):]
                    anchor = last_bullet_p
                    for extra_text in extras_available[:extras_needed]:
                        anchor = _clone_paragraph_after(anchor, extra_text)

                rewritten_count += 1
        current_bullet_indices = []

    while i < end_idx:
        p = paragraphs[i]
        if _paragraph_is_blank(p):
            i += 1
            continue
        if _is_role_header(p):
            _flush()
            current_company, current_title = _extract_company_and_title(p.text)
        else:
            if current_company:
                current_bullet_indices.append(i)
        i += 1
    _flush()

    return rewritten_count


def _normalize_stray_bullet_glyphs(doc) -> None:
    """Final pass: any paragraph that starts with `●` (or other heavy-bullet
    glyphs) gets normalized to `•` so the generated resume uses a single glyph
    throughout. Leaves non-bullet paragraphs untouched."""
    stray_prefix_re = re.compile(r"^([\u25cf\u25a0\u25aa\u2043])(\s*)")
    for p in doc.paragraphs:
        if not p.runs:
            continue
        first_run = p.runs[0]
        text = first_run.text or ""
        m = stray_prefix_re.match(text)
        if not m:
            continue
        first_run.text = "\u2022" + m.group(2) + text[m.end():]


def _remove_duplicate_trailing_blocks(doc) -> None:
    """Some source templates accidentally ship a stray copy of an earlier
    paragraph near the end of the document (manual copy/paste in Word). Strip
    any paragraph whose stripped text is identical to an earlier non-empty
    paragraph when it appears AFTER the Education heading."""
    seen_before_end: set[str] = set()
    education_seen = False
    to_clear = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        low = text.lower()
        if not education_seen:
            if _MARKER_EDUCATION in low and low.startswith(">"):
                education_seen = True
            seen_before_end.add(text)
            continue
        if text in seen_before_end:
            to_clear.append(p)
    for p in to_clear:
        for run in p.runs:
            run.text = ""


def generate_tailored_resume_bytes(
    archetype_id: str,
    job_description: str,
    use_llm: bool = False,
) -> bytes:
    """Produce the tailored DOCX as bytes, ready to send over HTTP."""
    template_path = _resolve_template_path(archetype_id)
    doc = Document(str(template_path))

    paragraphs = list(doc.paragraphs)
    idx = _find_indices(paragraphs)

    summary_text = draft_summary(job_description, archetype_id)
    if use_llm and _llm_available():
        summary_text = _llm_rewrite_summary(summary_text, job_description, archetype_id)
    _rewrite_summary(paragraphs, summary_text, idx)
    _rewrite_experience(paragraphs, idx, job_description, use_llm=use_llm)
    _normalize_stray_bullet_glyphs(doc)
    _remove_duplicate_trailing_blocks(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


__all__ = ["generate_tailored_resume_bytes"]
